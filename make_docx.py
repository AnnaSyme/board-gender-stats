"""
Generate README.docx from README content.
Uses python-docx; images are embedded at A4 content width.
"""
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

BASE   = "/Users/symea1/ClaudeBoardStats"
DATA   = os.path.join(BASE, "data")
OUTPUT = os.path.join(BASE, "README.docx")

MAX_W = Cm(17)
MAX_H = Cm(18)


def set_font(run, bold=False, code=False):
    run.bold = bold
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    else:
        run.font.name = "Helvetica Neue"
        run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x44)


def add_para(doc, text, bold=False, code=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_font(run, bold=bold, code=code)
    return p


def add_mixed(doc, parts, space_after=4):
    """parts = list of (text, bold, code) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, bold, code in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, code=code)
    return p


def add_image(doc, filename, space_after=8):
    path = os.path.join(DATA, filename)
    if not os.path.exists(path):
        return
    with PILImage.open(path) as im:
        pw, ph = im.size
    aspect = ph / pw
    w = MAX_W
    h = w * aspect
    if h > MAX_H:
        h = MAX_H
        w = h / aspect
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=int(w), height=int(h))


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x44)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F8")
        pPr.append(shd)


doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(section, attr, Cm(2))

# ── Content ───────────────────────────────────────────────────────────────────

add_para(doc, 'Dr Seuss famously wrote about \u201cToo many Daves\u201d. But are there too many Daves on the board? ')
add_para(doc, 'Let\u2019s have a look.')

# Chart 1
add_para(doc, '1. Number of boards with more Daves than women', bold=True)
add_para(doc, 'Around the table of 126 ASX boards, there are more Davids than there are women.')
add_para(doc, 'At 111 tables there are more Peters; at 104 there are more Andrews; at 103 there are more Michaels.')
add_image(doc, 'names_comparison_gender.png')

# Chart 2
add_para(doc, '2. Number of boards with more Kates than men', bold=True)
add_para(doc, 'Are there also too many Kates? Janes? Michelles?\u202f')
add_para(doc, 'No.\u202f')
add_image(doc, 'names_comparison_gender2.png')

# Chart 3
add_para(doc, '3. Who\u2019s at the table', bold=True)
add_para(doc, 'Did you know that 12% of ASX boards have no women on them? Just joking; it\u2019s 51%.')
add_para(doc, 'Most boards are mostly men:')
add_image(doc, 'chart_boardroom_two.png')

# Chart 4
add_para(doc, '4. Not all men', bold=True)
add_para(doc, 'But not all boards are all men:')
add_image(doc, 'chart_boardroom.png')

# Chart 5
add_para(doc, '5. Who are we shaking hands with', bold=True)
add_para(doc, 'David, Michael, Peter and Andrew. Matthew, Mark, Luke and John. Kate.')
add_image(doc, 'chart_top_names.png')

add_image(doc, 'chart_boardroom_names.png')

# Inspiration
add_mixed(doc, [
    ('Inspiration:', True, False),
    (' Deb Verhoeven\u2019s work on Daversity. '
     'https://www.abc.net.au/news/science/2017-11-24/australian-research-has-a-daversity-problem/9178786',
     False, False),
])

# Data
add_mixed(doc, [
    ('Data:', True, False),
    (' Board member data is fetched from the MarkitDigital API used by the ASX website (no API key required). '
     'Gender is inferred from name prefixes (Mr/Sir/Lord \u2192 male; Ms/Mrs/Miss/Dame \u2192 female). '
     'Titles like Dr or Prof. are classified as unknown (~4%). ', False, False),
    ('data/directors.csv', False, True),
    (' \u2014 one row per board seat: ', False, False),
    ('ticker', False, True), (', ', False, False),
    ('company', False, True), (', ', False, False),
    ('raw_name', False, True), (', ', False, False),
    ('clean_name', False, True), (', ', False, False),
    ('first_name', False, True), (', ', False, False),
    ('title', False, True), (', ', False, False),
    ('gender', False, True), (', ', False, False),
    ('is_board', False, True), ('.', False, False),
])

# Usage
add_mixed(doc, [
    ('Usage:', True, False),
    (' Requires Python 3 and ', False, False),
    ('matplotlib', False, True),
    ('. Run the scripts in this order:', False, False),
])
add_code_block(doc, [
    'python3 collect_boards.py          # fetch board data from the ASX/MarkitDigital API (~20 min); saves data/directors.csv',
    'python3 chart_names_comparison.py  # charts 1 & 2 \u2192 data/names_comparison_gender.png, names_comparison_gender2.png',
    'python3 chart_boardroom.py         # charts 3 & 4 \u2192 data/chart_boardroom_two.png, chart_boardroom.png',
    'python3 chart_top_names.py         # chart 5 \u2192 data/chart_top_names.png',
    'python3 chart_gospel_women.py      # charts 6 & 7 \u2192 data/chart_gospel_men.png, chart_gospel_women.png',
    'python3 chart_boardroom_names.py   # bonus: named seats chart \u2192 data/chart_boardroom_names.png',
    'python3 chart_boardroom_groups.py  # bonus: grouped seats chart \u2192 data/chart_boardroom_groups.png',
])

# Name combinations
add_mixed(doc, [
    ('Name combinations:', True, False),
    (' Some name variants are combined and counted together. '
     'Men: David+Dave, Michael+Mike+Mick, John+Jon, Stephen+Steven+Steve, James+Jim+Jamie, '
     'Christopher+Chris, Matthew+Matt, Timothy+Tim, Philip+Phillip+Phil, Peter+Pete, Andrew+Andy, '
     'Gregory+Greg, Geoffrey+Geoff. '
     'Women: Kate+Katherine+Kathryn+Kathy+Katie, Sarah+Sara, Anne+Anna+Ann+Annie, '
     'Jennifer+Jenny, Christine+Christina, Susan+Sue.', False, False),
])

# Licence
add_mixed(doc, [
    ('Licence:', True, False),
    (' Code is released under the MIT Licence. Charts are released under CC BY 4.0 \u2014 '
     'you may use and share them with attribution. '
     'The text of this README is \u00a9 Anna Syme, all rights reserved.', False, False),
])

# Built by
add_mixed(doc, [
    ('Built by', True, False),
    (' Anna Syme and Claude (Anthropic).', False, False),
])

doc.save(OUTPUT)
print(f"Saved {OUTPUT}")
